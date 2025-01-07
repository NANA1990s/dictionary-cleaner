import sys
import re
import os
from docx import Document
from pathlib import Path
from .cleaner import DictionaryCleaner

def get_user_choice(file_path: str) -> tuple[bool, str]:
    """获取用户选择
    
    Returns:
        tuple: (是否继续, 输出文件路径)
    """
    while True:
        print(f"\n发现已存在的输出文件：{file_path}")
        print("\n请选择操作：")
        print("1. 替换已有文件")
        print("2. 使用新文件名")
        print("3. 取消操作")
        
        choice = input("\n请输入选项（1-3）：").strip()
        
        if choice == "1":
            return True, file_path
        elif choice == "2":
            while True:
                new_name = input("\n请输入新的文件名（不包含路径和扩展名）：").strip()
                if new_name:
                    new_path = str(Path(file_path).parent / f"{new_name}.docx")
                    if not os.path.exists(new_path):
                        return True, new_path
                    print(f"文件 {new_path} 已存在，请使用其他名称。")
        elif choice == "3":
            return False, ""
        
        print("无效的选择，请重试。")

def clean_word_document(input_path: str) -> None:
    """处理 Word 文档，清理每个段落中的例句"""
    try:
        # 创建清理器
        cleaner = DictionaryCleaner()
        
        # 读取文档
        doc = Document(input_path)
        
        # 生成默认输出文件名
        output_path = str(Path(input_path).with_stem(Path(input_path).stem + '_cleaned'))
        
        # 如果文件已存在，获取用户选择
        if os.path.exists(output_path):
            should_continue, output_path = get_user_choice(output_path)
            if not should_continue:
                print("操作已取消")
                return
        
        # 收集所有段落文本
        text_blocks = []
        current_block = []
        empty_line_positions = []  # 记录空行位置
        current_position = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:  # 空行
                if current_block:  # 如果当前块不为空，记录空行位置
                    empty_line_positions.append(current_position)
            else:
                # 如果是新词条（以字母开头）或者是标记行，开始新的块
                if (re.match(r'^[A-Za-z]', text) and not cleaner.is_english_example(text)) or \
                   text.startswith(('[', '【', '(', '（', '||')) or \
                   re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text):
                    if current_block:
                        text_blocks.append('\n'.join(current_block))
                        current_block = []
                    current_block.append(text)
                else:
                    # 如果不是英文例句，添加到当前块
                    if not cleaner.is_english_example(text):
                        current_block.append(text)
            current_position += 1
        
        # 添加最后一个块
        if current_block:
            text_blocks.append('\n'.join(current_block))
        
        # 清理文档
        new_doc = Document()
        current_position = 0
        
        # 处理每个文本块
        for block in text_blocks:
            cleaned_text = cleaner.clean_sentence(block)
            # 按换行分割并添加到文档
            for line in cleaned_text.split('\n'):
                if line.strip():
                    new_doc.add_paragraph(line.strip())
                    current_position += 1
                    # 检查是否需要添加空行
                    if current_position in empty_line_positions:
                        new_doc.add_paragraph()
                        current_position += 1
        
        # 保存文档
        new_doc.save(output_path)
        print(f"\n处理完成！输出文件：{output_path}")
        
    except Exception as e:
        raise Exception(f"处理文件时出错：{str(e)}")

def main():
    """主函数：处理命令行参数并执行文档清理"""
    if len(sys.argv) != 2:
        print("使用方法: python clean_dictionary.py <input_file.docx>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    
    # 检查文件格式
    if not input_file.endswith('.docx'):
        print("错误：请提供 .docx 格式的 Word 文档")
        sys.exit(1)
    
    # 检查文件是否存在
    if not Path(input_file).exists():
        print(f"错误：文件 {input_file} 不存在")
        sys.exit(1)
    
    try:
        clean_word_document(input_file)
    except Exception as e:
        print(f"错误：{str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()