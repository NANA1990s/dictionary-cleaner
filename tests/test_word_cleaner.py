# Copyright 2024 [NANA1990s]
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import sys
import os
import pytest
from docx import Document
from pathlib import Path

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from src.clean_dictionary import clean_word_document

def test_word_document_cleaning(tmp_path):
    """测试 Word 文档处理功能"""
    # 创建测试文档
    doc = Document()
    test_cases = [
        "abandon v. 抛弃；放弃：The captain refused to abandon his sinking ship. 船长拒绝弃船。",
        "[近义词] desert, forsake",
        "[反义词] keep, maintain",
        "[谚语] Never abandon hope."
    ]
    
    for text in test_cases:
        doc.add_paragraph(text)
    
    # 保存测试文档
    test_file = tmp_path / "test.docx"
    doc.save(test_file)
    
    # 处理文档
    clean_word_document(str(test_file))
    
    # 检查结果
    output_file = tmp_path / "test_cleaned.docx"
    assert output_file.exists()
    
    # 验证内容
    result_doc = Document(output_file)
    expected = [
        "abandon v. 抛弃；放弃",
        "[近义词] desert, forsake",
        "[反义词] keep, maintain",
        "[谚语] Never abandon hope."
    ]
    
    for p, exp in zip(result_doc.paragraphs, expected):
        assert p.text.strip() == exp
